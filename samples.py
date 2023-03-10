# scratchpad
=================
import pandas as pd
import json

# Load Excel file into a pandas DataFrame
df = pd.read_excel('file.xlsx')

# Load JSON schema into a pandas DataFrame
with open('schema.json') as json_file:
    schema = json.load(json_file)
    schema_df = pd.DataFrame(schema, columns=['l3', 'other_columns'])

# Merge the two DataFrames based on the l3 column with outer join
merged_df = pd.merge(df, schema_df, on='l3', how='outer')
====================================

import pandas as pd
import json

# Load Excel file into a pandas DataFrame
df = pd.read_excel('file.xlsx')

# Load JSON schema into a pandas DataFrame
with open('schema.json') as json_file:
    schema = json.load(json_file)
    schema_df = pd.DataFrame(schema, columns=['l3', 'other_columns'])

# Merge the two DataFrames based on the l3 column
merged_df = pd.merge(df, schema_df, on='l3')

========================================
import pandas as pd

# Load data into a pandas DataFrame
df = pd.read_excel('file.xlsx')

# Get rows with NaN values
rows_with_NaN = df[df.isna().any(axis=1)]
============merge two dataframes ===============
import pandas as pd

# Load the data into two DataFrames
df1 = pd.read_excel("df1.xlsx")
df2 = pd.read_excel("df2.xlsx")

# Merge the two DataFrames on a common column
df_merged = df1.merge(df2[['common_column_name', 'desired_column_name']], on='common_column_name', how='left')
# Sort the DataFrame by the date column in ascending order
df = df.sort_values(by='date_column', ascending=True)
==============================================
# Iterate over the rows of the DataFrame
for index, row in df.iterrows():
    # Access data for each column by column name
    print('Column1:', row['Column1'], 'Column2:', row['Column2'])
====================================================
import pandas as pd

# Create a sample DataFrame
df = pd.DataFrame({'Column1': [1, 2, 3, 4], 'Column2': ['A', 'B', 'C', 'D']})

# Select the 'Column1' column
column1 = df['Column1']

# Convert the column to a list
column1_list = column1.tolist()

print(column1_list)
============= merege df and dict ================
import pandas as pd

# Create the dataframe
df = pd.DataFrame({'key': ['A', 'B', 'C', 'D'], 'value_1': [1, 2, 3, 4]})

# Create the dictionary
d = {'key': 'B', 'value_2': 10}

# Convert the dictionary to a dataframe
d = pd.DataFrame(d, index=[0])

# Merge the two dataframes based on the 'key' column
df = pd.merge(df, d, on='key', how='left')

print(df)
=============================================
for index, row in df.iterrows():
    r = requests.get(row['ListOfURLs'])
    if r.status_code == 200:
        df.at[index, ['Status Code', 'Result', 'Error']] = (r.status_code, '[OK]', np.nan)

print(df)
================================================
def get_algo_mcs(self):
    df_algo_names = df_algo_names[df_algo_names['Algo Name'].apply(lambda x: isinstance(x, str))]
    df_algo_names = df_algo_names[df_algo_names['Algo Name'].apply(lambda x: len(x) >= 5)]

    for index, row in df_algo_names.iterrows():
        algo = row['Algo Name']
        algoid, status = self._search(algo)
        if algoid != "ERR":
            df_algo_names.at[index, 'mode'] = "modelunumer"
====================================================            
lgEnt = [lg.get("code") for lg in md_data["info"]["legalENT"] if lg.get("code") is not None]
isEASY = 'Y' if 'EASY' in lgEnt else 'N'
===============================================
original_dict = {'a': 1, 'b': 2, 'c': 3, 'd': 4}
even_dict = {key: value for key, value in original_dict.items() if value % 2 == 0}
print(even_dict)
# Output: {'b': 2, 'd': 4}
=================compare DFs==============================
import pandas as pd

df1 = pd.DataFrame({
    'A': ['A0', 'A1', 'A2', 'A3'],
    'B': ['B0', 'B1', 'B2', 'B3'],
    'C': ['C0', 'C1', 'C2', 'C3'],
    'D': ['D0', 'D1', 'D2', 'D3']
})

df2 = pd.DataFrame({
    'A': ['A2', 'A3', 'A4', 'A5'],
    'B': ['B2', 'B3', 'B4', 'B5'],
    'C': ['C2', 'C3', 'C4', 'C5'],
    'D': ['D2', 'D3', 'D4', 'D5']
})

merged_left = df1.merge(df2, on=['A', 'B'], how='left')
merged_right = df1.merge(df2, on=['A', 'B'], how='right')
merged_outer = df1.merge(df2, on=['A', 'B'], how='outer')
merged_df = pd.merge(df1, df2, left_on='key1', right_on='key2')

import pandas as pd

# create the first DataFrame
df1 = pd.DataFrame({'key1': [1, 2, 3, 4], 'key2': ['A', 'B', 'C', 'D'], 'value1': ['X', 'Y', 'Z', 'W']})

# create the second DataFrame
df2 = pd.DataFrame({'key3': [3, 4, 5, 6], 'key4': pd.to_datetime(['2020-01-01', '2020-01-02', '2020-01-03', '2020-01-04']), 'value2': ['P', 'Q', 'R', 'S']})

# convert key2 to datetime
df1['key2'] = pd.to_datetime(df1['key2'])

# merge the two DataFrames based on the key1 and key3 columns, and key2 and key4 columns
merged_df = pd.merge(df1, df2, left_on=['key1', 'key2'], right_on=['key3', 'key4'])
if isinstance(df['col3'].dtype, pd.DatetimeTZDtype):
# display the merged DataFrame
print(merged_df)
import pandas as pd

# create a DataFrame
df = pd.DataFrame({'col1': [1, 2, 3, 4], 'col2': ['A', 'B', 'C', 'D'], 'col3': ['2020-01-01', '2020-01-02', '2020-01-03', '2020-01-04']})

# convert col3 to datetime and then to string with format 'dd-mm-yy'
df['col3'] = pd.to_datetime(df['col3'])
df['col3'] = df['col3'].dt.strftime('%d-%m-%y')

# display the resulting DataFrame
print(df)
df_algo_names['ATGF Date'] = pd.to_datetime(df_algo_names['ATGF Date'].astype(str), format='%Y-%m-%d %H:%M:%S')
==========================================
import os
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE

# Open the DOCX file
docx_file = 'example.docx'
doc = Document(docx_file)

# Get the attachment relationships
attachment_rels = [
    rel for rel in doc.part.rels.values() 
    if rel.reltype == RELATIONSHIP_TYPE.ATTACHMENT
]

# Extract the attachments
for rel in attachment_rels:
    # Get the attachment name and path
    attachment_name = os.path.basename(rel.target_ref.path)
    attachment_path = os.path.join(
        os.path.dirname(docx_file), rel.target_ref.path
    )
    # Save the attachment to disk
    with open(attachment_path, 'wb') as f:
        f.write(rel.target_part.blob)
    print(f'Extracted attachment: {attachment_name}')
====================================================
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE

# Open the DOCX file
docx_file = 'example.docx'
doc = Document(docx_file)

# Get the attachment relationships
attachment_rels = [
    rel for rel in doc.part.rels.values() 
    if rel.reltype == RELATIONSHIP_TYPE.ATTACHMENT
]

# Extract the PDF attachments
for rel in attachment_rels:
    if rel.target_part.content_type == 'application/pdf':
        # Get the attachment name and title
        attachment_name = rel.target_ref.path
        attachment_title = rel.target_part.get_or_add_title_of_image_part()

        # Save the attachment to disk
        with open(attachment_name, 'wb') as f:
            f.write(rel.target_part.blob)
        
        print(f'Extracted attachment: {attachment_title}')
=================================
from docx import Document

# Open the DOCX file
docx_file = 'example.docx'
doc = Document(docx_file)

# Get the attachment relationships
attachment_rels = [
    rel for rel in doc.part.rels.values() 
    if rel.target_part.content_type == 'application/pdf'
]

# Extract the PDF attachments
for rel in attachment_rels:
    # Get the attachment name and path
    attachment_name = rel.target_ref.path
    attachment_path = os.path.join(
        os.path.dirname(docx_file), rel.target_ref.path
    )
    # Save the attachment to disk
    with open(attachment_path, 'wb') as f:
        f.write(rel.target_part.blob)
    print(f'Extracted attachment: {attachment_name}')
==========================
from docx import Document
import requests

# Open the DOCX file
docx_file = 'example.docx'
doc = Document(docx_file)

# Get the attachment relationships
attachment_rels = [
    rel for rel in doc.part.rels.values()
    if rel.target_part is None and rel.is_external
]

# Extract the external attachments
for rel in attachment_rels:
    # Get the attachment name and target URI
    attachment_name = rel.target_ref.path
    attachment_uri = rel.target_ref.uri

    # Download the attachment from the target URI
    response = requests.get(attachment_uri)

    # Save the attachment to disk
    with open(attachment_name, 'wb') as f:
        f.write(response.content)

    print(f'Extracted attachment: {attachment_name}')
==============================================
import os
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE

# Open the DOCX file
docx_file = 'example.docx'
doc = Document(docx_file)

# Find the attachment relationship with TargetMode='External'
attachment_rel = None
for rel in doc.part.rels.values():
    if rel.rel_type == RELATIONSHIP_TYPE.HYPERLINK:
        # check if the target is an external file
        if rel.target_mode == RELATIONSHIP_TARGET_MODE_EXTERNAL:
            attachment_rel = rel
            break

if attachment_rel is None:
    print('No external attachment found')
else:
    # Get the original file name from the relationship
    attachment_name = os.path.basename(attachment_rel.target_ref.target)
    
    # Extract the attachment contents to a file with the original name
    with open(attachment_name, 'wb') as f:
        f.write(attachment_rel.target_ref.part.blob)

    print(f'Extracted attachment: {attachment_name}')
